from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
from typing import Optional
from services.chargesheet_service import ChargesheetService

router = APIRouter(
    prefix="/api/chargesheet-generation",
    tags=["Chargesheet Generation"]
)

# Initialize Service
service = ChargesheetService()

@router.post("")
async def generate_chargesheet(
    fir_document: Optional[UploadFile] = File(None),
    case_id: Optional[str] = Form(None),
    incident_details_file: Optional[UploadFile] = File(None),
    incident_details_text: Optional[str] = Form(None),
    additional_instructions: Optional[str] = Form(None)
):
    try:
        # Validate Inputs: Need either File OR Case ID
        if not fir_document and not case_id:
            raise HTTPException(status_code=400, detail="Either FIR Document upload or Case ID is required.")

        # 1. Extract content
        fir_text = None
        if fir_document:
            fir_text = await service.extract_text_from_file(fir_document)
            
        incident_text = ""
        if incident_details_file:
            incident_text += await service.extract_text_from_file(incident_details_file)
        
        if incident_details_text and incident_details_text.strip():
            if incident_text:
                incident_text += "\n\n"
            incident_text += incident_details_text.strip()
            
        # 2. Call Service
        draft = await service.generate_draft(
            fir_text=fir_text,
            case_id=case_id,
            incident_text=incident_text,
            instructions=additional_instructions
        )
        
        return JSONResponse(content={"chargeSheet": draft})
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

from fastapi.responses import StreamingResponse
import io
from docx import Document
from docx.shared import Pt

@router.post("/download-docx")
async def download_chargesheet_docx(
    chargesheetText: str = Form(...)
):
    try:
        doc = Document()
        # Add styled Heading
        heading = doc.add_heading('CHARGE SHEET', 0)
        heading.alignment = 1  # Center alignment
        run = heading.runs[0]
        run.bold = True
        run.font.size = Pt(16)  # 16pt bold
        
        # Split by newlines and add paragraphs
        for line in chargesheetText.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph("") # maintain spacing
            
        byte_io = io.BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        
        return StreamingResponse(
            byte_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=chargesheet.docx"}
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")
